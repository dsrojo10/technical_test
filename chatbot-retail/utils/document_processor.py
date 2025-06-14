import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from typing import List, Dict
import logging
from pathlib import Path


class DocumentProcessor:
    """Procesador de documentos para extraer texto de diferentes formatos"""
    
    @staticmethod
    def extract_from_xlsx(file_path: str) -> str:
        """
        Extrae texto de un archivo Excel
        
        Args:
            file_path: Ruta al archivo Excel
            
        Returns:
            str: Texto extraído del archivo
        """
        try:
            if not Path(file_path).exists():
                logging.warning(f"Archivo no encontrado: {file_path}")
                return ""
            
            # Leer todas las hojas del Excel
            excel_file = pd.ExcelFile(file_path)
            text_content = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Agregar nombre de la hoja
                text_content.append(f"HOJA: {sheet_name}")
                
                # Convertir DataFrame a texto
                # Incluir nombres de columnas
                if not df.empty:
                    text_content.append(df.to_string(index=False))
                
                text_content.append("\n" + "="*50 + "\n")
            
            return "\n".join(text_content)
            
        except Exception as e:
            logging.error(f"Error procesando Excel {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """
        Extrae texto de un archivo PDF
        
        Args:
            file_path: Ruta al archivo PDF
            
        Returns:
            str: Texto extraído del archivo
        """
        try:
            if not Path(file_path).exists():
                logging.warning(f"Archivo no encontrado: {file_path}")
                return ""
            
            reader = PdfReader(file_path)
            text_content = []
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"PÁGINA {i + 1}:")
                    text_content.append(text)
                    text_content.append("\n" + "="*50 + "\n")
            
            return "\n".join(text_content)
            
        except Exception as e:
            logging.error(f"Error procesando PDF {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """
        Extrae texto de un archivo Word
        
        Args:
            file_path: Ruta al archivo Word
            
        Returns:
            str: Texto extraído del archivo
        """
        try:
            if not Path(file_path).exists():
                logging.warning(f"Archivo no encontrado: {file_path}")
                return ""
            
            doc = Document(file_path)
            text_content = []
            
            # Extraer párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extraer texto de tablas si las hay
            for table in doc.tables:
                text_content.append("\nTABLA:")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
                text_content.append("="*50)
            
            return "\n".join(text_content)
            
        except Exception as e:
            logging.error(f"Error procesando DOCX {file_path}: {str(e)}")
            return ""
    
    @classmethod
    def process_all_documents(cls, document_paths: Dict[str, str]) -> Dict[str, str]:
        """
        Procesa múltiples documentos y retorna el texto extraído
        
        Args:
            document_paths: Diccionario con nombre_documento: ruta_archivo
            
        Returns:
            Dict con nombre_documento: texto_extraído
        """
        processed_docs = {}
        
        for doc_name, file_path in document_paths.items():
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.xlsx':
                text = cls.extract_from_xlsx(file_path)
            elif file_extension == '.pdf':
                text = cls.extract_from_pdf(file_path)
            elif file_extension == '.docx':
                text = cls.extract_from_docx(file_path)
            else:
                logging.warning(f"Formato no soportado para {file_path}")
                text = ""
            
            processed_docs[doc_name] = text
        
        return processed_docs
    
    @staticmethod
    def split_text_into_chunks(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Divide el texto en chunks para embeddings
        
        Args:
            text: Texto a dividir
            chunk_size: Tamaño máximo de cada chunk
            overlap: Superposición entre chunks
            
        Returns:
            Lista de chunks de texto
        """
        if not text or len(text) <= chunk_size:
            return [text] if text else []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Si no es el último chunk, buscar un punto de corte natural
            if end < len(text):
                # Buscar el último espacio o salto de línea antes del límite
                last_space = text.rfind(' ', start, end)
                last_newline = text.rfind('\n', start, end)
                
                # Usar el corte más tardío que sea natural
                natural_cut = max(last_space, last_newline)
                if natural_cut > start:
                    end = natural_cut
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Mover el inicio con overlap
            start = end - overlap if end < len(text) else end
        
        return chunks
